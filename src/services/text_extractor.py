"""
Text Extractor Service - Document Processor

Extract text from PDF and DOCX files with automatic normalization.

Hybrid PDF Extraction Strategy:
    - PDF (strategy='fast'): pdfplumber - Fast, simple layouts, default
    - PDF (strategy='deep'): marker-pdf - Slow, complex multi-column layouts, AI-powered
    - DOCX: python-docx - Always used (ignores strategy parameter)

The hybrid strategy allows users to choose between:
    1. Fast extraction (pdfplumber): Good for most PDFs, single or simple multi-column
    2. Deep extraction (marker-pdf): Required for complex layouts like Lexis PDFs
       where headers and bodies have different column counts on same page

Future strategies can be added:
    - 'medical': Optimize for medical documents
    - 'ocr_only': Force OCR even on text PDFs
    - 'table_mode': Preserve table structures

Returns:
    ExtractionResult model with normalized, clean text

Usage:
    from src.services.text_extractor import extract_text

    # Standard extraction (fast)
    result = extract_text(Path("/files/document.pdf"))
    result = extract_text(Path("/files/document.pdf"), strategy='fast')

    # Deep extraction (complex layouts)
    result = extract_text(Path("/files/lexis.pdf"), strategy='deep')

    if result.success:
        print(result.text)
    else:
        print(result.error_message)
"""

from pathlib import Path
from typing import Optional, Tuple

from src.core.models import ExtractionResult
from src.cleaners.text_normalizer import normalize_text


# ============================================================================
# EXTRACTION STRATEGIES (Private Functions)
# ============================================================================

def _extract_pdf_fast(file_path: Path, max_pages: Optional[int] = None) -> Tuple[str, int]:
    """
    Extract text from PDF using pdfplumber (fast mode).

    This is the default PDF extraction strategy for simple layouts.
    Uses pdfplumber's layout-preserving extraction which works well
    for single-column documents and simple multi-column layouts.

    Args:
        file_path: Path to PDF file
        max_pages: Optional limit on number of pages to extract

    Returns:
        Tuple of (extracted_text, page_count)

    Raises:
        Exception: Any pdfplumber error (corrupt PDF, unsupported format, etc.)

    Note:
        This is an internal strategy function. Use extract_text() instead.
        For complex multi-column layouts (e.g., Lexis PDFs), use _extract_pdf_deep().
    """
    import pdfplumber

    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)

        # Determine which pages to extract
        if max_pages is not None:
            pages_to_extract = pdf.pages[:max_pages]
        else:
            pages_to_extract = pdf.pages

        # Extract text from each page
        # Use layout=True to preserve column structure
        extracted_lines = []
        for page in pages_to_extract:
            # Try layout-preserving extraction first
            page_text = page.extract_text(layout=True)

            # Fallback to simple extraction if layout fails
            if not page_text:
                page_text = page.extract_text(layout=False)

            # If still no text, this might be an image-based PDF page
            if page_text:
                extracted_lines.append(page_text)

        # Join pages with clear separator
        full_text = "\n\n".join(extracted_lines)

        return full_text, total_pages


def _extract_docx(file_path: Path) -> Tuple[str, int]:
    """
    Extract text from DOCX using python-docx.

    Args:
        file_path: Path to DOCX file

    Returns:
        Tuple of (extracted_text, paragraph_count)

    Raises:
        Exception: Any python-docx error (corrupt file, unsupported format, etc.)

    Note:
        This is an internal strategy function. Use extract_text() instead.
    """
    from docx import Document

    doc = Document(file_path)

    # Extract text from all paragraphs
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also extract text from tables (common in legal documents)
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_texts.append(row_text)

    # Combine paragraphs and tables
    all_text = paragraphs + table_texts
    full_text = "\n\n".join(all_text)

    # Paragraph count as proxy for "page count" (DOCX doesn't have pages concept)
    paragraph_count = len(paragraphs)

    return full_text, paragraph_count


def _extract_legacy_doc(file_path: Path) -> Tuple[str, int]:
    """
    Extract text from legacy .doc files using LibreOffice conversion.

    Args:
        file_path: Path to .doc file

    Returns:
        Tuple of (extracted_text, page_count_estimate)

    Raises:
        Exception: If LibreOffice is not installed or conversion fails

    Note:
        This requires LibreOffice to be installed on the system.
        Converts .doc -> .docx -> extract text.
        This is slower than direct extraction.
    """
    import subprocess
    import tempfile

    # Create temporary directory for conversion
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)

        # Convert .doc to .docx using LibreOffice
        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(temp_path),
                    str(file_path),
                ],
                check=True,
                capture_output=True,
                timeout=60,  # 60 second timeout
            )
        except FileNotFoundError:
            raise Exception(
                "LibreOffice not found. Install LibreOffice to process legacy .doc files. "
                "On macOS: brew install --cask libreoffice"
            )
        except subprocess.TimeoutExpired:
            raise Exception(f"LibreOffice conversion timed out for {file_path.name}")

        # Find the converted file
        converted_file = temp_path / f"{file_path.stem}.docx"
        if not converted_file.exists():
            raise Exception(f"LibreOffice conversion failed for {file_path.name}")

        # Extract text from converted DOCX
        return _extract_docx(converted_file)


def _markdown_to_plain_text(markdown_text: str) -> str:
    """
    Convert markdown to plain text by stripping markdown syntax.

    This is used to convert marker-pdf's markdown output to plain text
    for consistency with pdfplumber's output format.

    Args:
        markdown_text: Markdown-formatted text from marker-pdf

    Returns:
        Plain text with markdown syntax removed

    Note:
        This performs basic markdown stripping (headers, bold, italic, links).
        Advanced markdown features may not be fully stripped.
    """
    import re

    # Remove markdown headers (### Header -> Header)
    text = re.sub(r'^#{1,6}\s+', '', markdown_text, flags=re.MULTILINE)

    # Remove bold/italic (**bold** or __bold__ -> bold)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)

    # Remove links ([text](url) -> text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    # Remove inline code (`code` -> code)
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # Remove code blocks (```...``` -> ...)
    text = re.sub(r'```[^\n]*\n(.*?)\n```', r'\1', text, flags=re.DOTALL)

    # Remove horizontal rules (--- or ***)
    text = re.sub(r'^(\*\*\*|---|___)\s*$', '', text, flags=re.MULTILINE)

    # Remove blockquotes (> text -> text)
    text = re.sub(r'^>\s+', '', text, flags=re.MULTILINE)

    # Remove list markers (- item or * item or 1. item -> item)
    text = re.sub(r'^[\*\-\+]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)

    return text


def _extract_pdf_deep(file_path: Path, max_pages: Optional[int] = None) -> Tuple[str, int]:
    """
    Extract text from PDF using marker-pdf (deep mode with AI).

    This uses AI-powered extraction for complex multi-column layouts
    that simple pdfplumber heuristics cannot handle correctly.
    Requires marker-pdf package (optional dependency, ~3-5GB with models).

    Args:
        file_path: Path to PDF file
        max_pages: Optional limit (NOTE: marker processes full PDF, truncation applied post-processing)

    Returns:
        Tuple of (extracted_text, page_count)

    Raises:
        ImportError: If marker-pdf is not installed
        Exception: Any marker-pdf error (corrupt PDF, unsupported format, etc.)

    Note:
        This is an internal strategy function. Use extract_text() instead.
        First time use will download ~3-5GB of PyTorch models.
        Installation: pip install marker-pdf

    Architecture:
        - Uses marker.converters.pdf.PdfConverter
        - Returns markdown format (converted to plain text)
        - Handles complex layouts: multi-column, tables, figures
    """
    # Lazy import marker-pdf (only when deep mode is used)
    try:
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        from marker.output import text_from_rendered
    except ImportError:
        raise ImportError(
            "marker-pdf is required for deep PDF extraction but not installed.\n"
            "Install with: pip install marker-pdf\n"
            "Note: This will download ~3-5GB of PyTorch models on first use.\n"
            "For standard extraction, use use_deep_extraction=False (default)."
        )

    # Initialize marker converter (lazy model loading)
    converter = PdfConverter(artifact_dict=create_model_dict())

    # Convert PDF to markdown
    rendered = converter(str(file_path))

    # Extract text from rendered result
    markdown_text, _, _ = text_from_rendered(rendered)

    # Convert markdown to plain text for consistency
    plain_text = _markdown_to_plain_text(markdown_text)

    # Get page count from rendered metadata
    page_count = len(rendered.pages) if hasattr(rendered, 'pages') else 0

    # Apply max_pages truncation if specified (post-processing)
    # Note: marker processes full PDF, so this is less efficient than pdfplumber's page limiting
    if max_pages is not None and page_count > max_pages:
        # Rough heuristic: split by double newlines and take proportional amount
        sections = plain_text.split('\n\n')
        target_sections = int(len(sections) * (max_pages / page_count))
        plain_text = '\n\n'.join(sections[:target_sections])

    return plain_text, page_count


# ============================================================================
# PUBLIC API
# ============================================================================

def extract_text(
    file_path: Path,
    max_pages: Optional[int] = None,
    normalize: bool = True,
    strategy: str = 'fast',
) -> ExtractionResult:
    """
    Extract text from PDF or DOCX file with automatic normalization.

    This is the main public API for text extraction. It:
        1. Detects file type (PDF, DOCX, DOC)
        2. Dispatches to appropriate extraction strategy
        3. Normalizes text (unicode, hyphens, whitespace)
        4. Returns clean ExtractionResult

    Hybrid PDF Extraction Strategy:
        - DOCX: ALWAYS uses python-docx (ignores strategy parameter)
        - PDF (strategy='fast'): Uses pdfplumber - fast, works for simple layouts
        - PDF (strategy='deep'): Uses marker-pdf AI - slow, handles complex multi-column layouts

    Args:
        file_path: Absolute path to document file
        max_pages: Optional limit on pages to extract (PDF only, less efficient with deep mode)
        normalize: Whether to normalize extracted text (default: True)
        strategy: Extraction strategy for PDFs (default: 'fast')
                 Options: 'fast' (pdfplumber) or 'deep' (marker-pdf AI)
                 Requires marker-pdf package (~3-5GB) for 'deep' mode
                 Ignored for DOCX files (always uses python-docx)

    Returns:
        ExtractionResult model with:
            - text: Extracted and normalized text
            - success: True if extraction succeeded
            - page_count: Number of pages/paragraphs
            - error_message: Error details if extraction failed

    Example:
        >>> # Standard extraction (fast, good for simple layouts)
        >>> result = extract_text(Path("/files/case.pdf"), strategy='fast')
        >>>
        >>> # Deep extraction (slow, handles complex multi-column layouts)
        >>> result = extract_text(Path("/files/lexis.pdf"), strategy='deep')
        >>>
        >>> # DOCX always uses python-docx (strategy parameter ignored)
        >>> result = extract_text(Path("/files/brief.docx"), strategy='deep')

    Error Handling:
        Returns ExtractionResult with success=False instead of raising exceptions.
        Possible errors:
            - FileNotFoundError: File doesn't exist
            - ImportError: marker-pdf not installed (when strategy='deep')
            - ValueError: Invalid strategy parameter
            - Corrupt PDF/DOCX
            - Unsupported file format
            - LibreOffice not installed (for .doc files)
            - Extraction timeout
    """
    # Validate file path
    if not file_path.exists():
        return ExtractionResult(
            text="",
            success=False,
            error_message=f"File not found: {file_path}",
        )

    if not file_path.is_file():
        return ExtractionResult(
            text="",
            success=False,
            error_message=f"Not a file: {file_path}",
        )

    # Validate strategy parameter
    valid_strategies = ['fast', 'deep']
    if strategy not in valid_strategies:
        return ExtractionResult(
            text="",
            success=False,
            error_message=f"Invalid strategy: '{strategy}'. "
            f"Supported strategies: {', '.join(valid_strategies)}",
        )

    # Detect file type and dispatch to strategy
    extension = file_path.suffix.lower()

    try:
        if extension == ".pdf":
            # PDF: Check extraction strategy (fast vs deep)
            if strategy == 'deep':
                raw_text, page_count = _extract_pdf_deep(file_path, max_pages)
            elif strategy == 'fast':
                raw_text, page_count = _extract_pdf_fast(file_path, max_pages)
            else:
                # This shouldn't happen due to validation above, but be defensive
                return ExtractionResult(
                    text="",
                    success=False,
                    error_message=f"Invalid PDF strategy: '{strategy}'",
                )

        elif extension == ".docx":
            # DOCX: ALWAYS use python-docx (ignore strategy parameter)
            raw_text, page_count = _extract_docx(file_path)

        elif extension == ".doc":
            # Legacy DOC: Convert to DOCX then extract
            raw_text, page_count = _extract_legacy_doc(file_path)

        else:
            return ExtractionResult(
                text="",
                success=False,
                error_message=f"Unsupported file format: {extension}. "
                f"Supported: .pdf, .docx, .doc",
            )

        # Check if extraction returned any text
        if not raw_text or not raw_text.strip():
            return ExtractionResult(
                text="",
                success=False,
                page_count=page_count,
                error_message=f"No text extracted from {file_path.name}. "
                f"This might be an image-based PDF (OCR required) or empty document.",
            )

        # Normalize text (unicode, hyphens, whitespace)
        if normalize:
            clean_text = normalize_text(raw_text)
        else:
            clean_text = raw_text

        # Return successful result
        return ExtractionResult(
            text=clean_text,
            success=True,
            page_count=page_count,
        )

    except Exception as e:
        # Catch all extraction errors and return graceful failure
        error_type = type(e).__name__
        return ExtractionResult(
            text="",
            success=False,
            error_message=f"{error_type}: {str(e)}",
        )


def extract_multizone(
    file_path: Path,
    first_pages: int = 5,
    last_pages: int = 2,
    normalize: bool = True,
) -> ExtractionResult:
    """
    Extract text from beginning and end of document (for metadata extraction).

    Useful for getting document metadata without processing entire file.
    Many documents have metadata at beginning (title page, headers)
    and end (signatures, dates, conclusions).

    Args:
        file_path: Absolute path to document file
        first_pages: Number of pages to extract from beginning (default: 5)
        last_pages: Number of pages to extract from end (default: 2)
        normalize: Whether to normalize extracted text (default: True)

    Returns:
        ExtractionResult with text from first and last pages.
        Text includes marker: "...[DOCUMENT MIDDLE OMITTED]..."

    Example:
        >>> result = extract_multizone(Path("/files/long_case.pdf"))
        >>> # Result contains first 5 + last 2 pages only
    """
    # Only implemented for PDFs currently
    if file_path.suffix.lower() != ".pdf":
        # For non-PDFs, just extract normally
        return extract_text(file_path, normalize=normalize)

    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)

            # Extract first pages
            first_text = "\n\n".join(
                page.extract_text() or ""
                for page in pdf.pages[:first_pages]
            )

            # Extract last pages (if document is long enough)
            if total_pages > first_pages + last_pages:
                last_text = "\n\n".join(
                    page.extract_text() or ""
                    for page in pdf.pages[-last_pages:]
                )

                # Combine with marker
                full_text = (
                    f"{first_text}\n\n"
                    f"...[DOCUMENT MIDDLE OMITTED: "
                    f"{total_pages - first_pages - last_pages} pages]...\n\n"
                    f"{last_text}"
                )
            else:
                # Document is short, just use first pages
                full_text = first_text

            # Normalize if requested
            if normalize:
                full_text = normalize_text(full_text)

            return ExtractionResult(
                text=full_text,
                success=True,
                page_count=total_pages,
            )

    except Exception as e:
        error_type = type(e).__name__
        return ExtractionResult(
            text="",
            success=False,
            error_message=f"{error_type}: {str(e)}",
        )


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def is_image_based_pdf(file_path: Path) -> bool:
    """
    Check if PDF is image-based (scanned) vs. text-based.

    Image-based PDFs require OCR to extract text.

    Args:
        file_path: Path to PDF file

    Returns:
        True if PDF appears to be image-based (no extractable text)

    Example:
        >>> if is_image_based_pdf(Path("scan.pdf")):
        ...     print("This PDF needs OCR")
    """
    try:
        result = extract_text(file_path, max_pages=3)

        if not result.success:
            # Can't determine, assume not image-based
            return False

        # If very little text extracted, probably image-based
        text_length = len(result.text.strip())
        return text_length < 50  # Less than 50 chars suggests scanned images

    except Exception:
        return False


def get_page_count(file_path: Path) -> Optional[int]:
    """
    Get page count without extracting full text.

    Args:
        file_path: Path to document file

    Returns:
        Page count, or None if unable to determine

    Example:
        >>> count = get_page_count(Path("case.pdf"))
        >>> print(f"Document has {count} pages")
    """
    extension = file_path.suffix.lower()

    try:
        if extension == ".pdf":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return len(pdf.pages)

        elif extension in [".docx", ".doc"]:
            from docx import Document
            doc = Document(file_path)
            return len(doc.paragraphs)  # Approximate

        else:
            return None

    except Exception:
        return None

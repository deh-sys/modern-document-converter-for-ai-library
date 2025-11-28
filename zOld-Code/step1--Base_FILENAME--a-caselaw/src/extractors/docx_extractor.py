"""
DOCX text extraction module.
Extracts text from Word documents for caselaw metadata processing.
"""

import re
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxExtractor:
    """Extract text from Word documents using python-docx."""

    def __init__(self, max_paragraphs=None):
        """
        Initialize DOCX extractor.

        Args:
            max_paragraphs: Maximum number of paragraphs to extract (default: None for all)
        """
        self.max_paragraphs = max_paragraphs

    @staticmethod
    def check_docx_available():
        """
        Check if python-docx library is available.

        Returns:
            tuple: (available: bool, error_message: str or None)
        """
        if DOCX_AVAILABLE:
            return (True, None)
        else:
            error_msg = (
                "python-docx not found. Please install it:\n"
                "  pip install python-docx\n"
                "  or: pip install -e ."
            )
            return (False, error_msg)

    def extract_text(self, docx_path):
        """
        Extract text from Word document.

        Args:
            docx_path: Path to .docx file

        Returns:
            str: Extracted text or None if extraction fails
        """
        if not DOCX_AVAILABLE:
            return None

        docx_path = Path(docx_path)

        if not docx_path.exists():
            return None

        try:
            doc = Document(str(docx_path))

            paragraphs = []
            para_count = 0

            # Extract text from all paragraphs
            for para in doc.paragraphs:
                if self.max_paragraphs and para_count >= self.max_paragraphs:
                    break

                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    para_count += 1

            # Extract text from tables
            for table in doc.tables:
                if self.max_paragraphs and para_count >= self.max_paragraphs:
                    break

                for row in table.rows:
                    for cell in row.cells:
                        if self.max_paragraphs and para_count >= self.max_paragraphs:
                            break

                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
                            para_count += 1

            return '\n'.join(paragraphs) if paragraphs else None

        except Exception as e:
            return None

    def get_main_content(self, text):
        """
        Filter out margin content and focus on main body text.
        Similar to PDF extractor's filtering logic.

        Args:
            text: Raw document text

        Returns:
            str: Filtered text with reduced margin content
        """
        if not text:
            return ""

        lines = text.split('\n')

        # Filter out very short lines that are likely headers/footers
        # But keep lines with key context words
        important_words = [
            'decided', 'filed', 'dated', 'court', 'circuit', 'district',
            'supreme', 'appeals', 'opinion', 'v.', 'vs.'
        ]

        filtered_lines = []
        for line in lines:
            line_stripped = line.strip()

            # Keep line if:
            # - It's substantial (>20 chars)
            # - OR it contains important context words
            if len(line_stripped) > 20 or any(
                word in line_stripped.lower()
                for word in important_words
            ):
                filtered_lines.append(line)

        return '\n'.join(filtered_lines)

    def extract_text_multizone(self, docx_path, first_paragraphs=50, last_paragraphs=20):
        """
        Extract text from multiple zones of document for comprehensive metadata extraction.
        Reads first N paragraphs + last M paragraphs to capture headers and conclusions.

        Args:
            docx_path: Path to .docx file
            first_paragraphs: Number of paragraphs from beginning (default: 50)
            last_paragraphs: Number of paragraphs from end (default: 20)

        Returns:
            str: Combined extracted text from both zones, or None if extraction fails
        """
        if not DOCX_AVAILABLE:
            return None

        docx_path = Path(docx_path)

        if not docx_path.exists():
            return None

        try:
            doc = Document(str(docx_path))

            # Collect all paragraphs (including from tables)
            all_paragraphs = []

            # Regular paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_paragraphs.append(text)

            # Table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            all_paragraphs.append(text)

            total_paragraphs = len(all_paragraphs)

            # Extract first paragraphs
            first_text = '\n'.join(all_paragraphs[:first_paragraphs])
            first_text = self.get_main_content(first_text)

            # Extract last paragraphs if document is long enough
            last_text = ''
            if total_paragraphs > first_paragraphs:
                last_start = max(first_paragraphs, total_paragraphs - last_paragraphs)
                last_text = '\n'.join(all_paragraphs[last_start:])
                last_text = self.get_main_content(last_text)

            # Combine both zones with a marker
            combined = first_text
            if last_text:
                combined += '\n\n[LAST_SECTION]\n\n' + last_text

            return combined if combined else None

        except Exception:
            # Fallback to regular extraction if multizone fails
            return self.extract_text(docx_path)

    def find_case_caption(self, text):
        """
        Find the case caption (case name) in document text.

        Args:
            text: Document text

        Returns:
            str: Case caption or None
        """
        if not text:
            return None

        lines = text.split('\n')

        # Look for "v." or "vs." pattern which indicates case name
        # Check first 40 lines
        for i, line in enumerate(lines[:40]):
            if re.search(r'\s+v\.?\s+', line, re.IGNORECASE):
                # Found a "v" - get context around it
                # Look at surrounding lines for full case name
                start = max(0, i - 2)
                end = min(len(lines), i + 3)
                context = '\n'.join(lines[start:end])

                # Try to extract the case name from this context
                # Look for pattern: "Party v. Party"
                match = re.search(
                    r'([A-Z][A-Za-z\s,\.&\'\-]+?)\s+v\.?\s+([A-Za-z\s,\.&\'\-]+?)(?:\n|$)',
                    context,
                    re.IGNORECASE
                )

                if match:
                    return f"{match.group(1).strip()} v. {match.group(2).strip()}"

        return None

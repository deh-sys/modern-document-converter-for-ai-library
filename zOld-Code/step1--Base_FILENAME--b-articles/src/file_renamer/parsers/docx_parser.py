"""DOCX parser with text extraction."""

import logging
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..models import DocumentMetadata, DocumentType
from .base import DocumentParser

LOGGER = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    """Parser for extracting text and metadata from Word documents."""

    document_type = DocumentType.DOCX

    def parse(self, path: Path) -> DocumentMetadata:
        """
        Extract metadata from a Word document.

        Currently extracts text but metadata extraction is not yet implemented.
        Returns basic metadata structure similar to PDF parser.
        """
        if not DOCX_AVAILABLE:
            LOGGER.warning("python-docx library not available, returning empty metadata")
            return DocumentMetadata(
                source_path=path,
                document_type=self.document_type,
            )

        # Extract text for future metadata extraction
        text = self._extract_text(path)

        # For now, return basic metadata
        # TODO: Integrate extraction logic to populate jurisdiction, decision_year, etc.
        return DocumentMetadata(
            source_path=path,
            document_type=self.document_type,
        )

    def _extract_text(self, path: Path) -> Optional[str]:
        """
        Extract all text content from a Word document.

        Args:
            path: Path to .docx file

        Returns:
            str: Extracted text content, or None if extraction fails
        """
        if not DOCX_AVAILABLE:
            return None

        try:
            doc = Document(str(path))

            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

            return '\n'.join(paragraphs)

        except Exception as e:
            LOGGER.exception("Failed to extract text from %s: %s", path, e)
            return None

    def supports(self, path: Path) -> bool:
        """Check if this parser can handle the given file."""
        return path.suffix.lower() == ".docx"
